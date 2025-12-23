#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 KET 单词默写练习 Word 文档
包含两种练习：
1. 根据汉语翻译填写英文单词
2. 根据英文和音标填写汉语翻译
"""

import json
import random
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def load_words(json_file):
    """从 JSON 文件加载单词列表"""
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    words_list = []
    for letter, words in data.items():
        for word_data in words:
            words_list.append({
                'word': word_data.get('word', ''),
                'phonetic': word_data.get('phonetic', ''),
                'chinese': word_data.get('chinese', '')
            })
    
    return words_list

def setup_document(doc):
    """设置文档页面格式"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

def format_cell(cell, font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """格式化单元格"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = font_size
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(14)
        paragraph.alignment = alignment

def generate_practice_chinese_to_english(doc, words_list, title="练习一：根据汉语翻译填写英文单词"):
    """生成练习一：根据汉语翻译填写英文单词"""
    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.size = Pt(14)
    
    # 添加说明
    instruction = doc.add_paragraph('说明：根据给出的汉语翻译，在横线上填写对应的英文单词。')
    instruction.runs[0].font.size = Pt(10)
    instruction.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()  # 空行
    
    # 创建表格（6列：汉语翻译1、英文单词1、音标1、汉语翻译2、英文单词2、音标2）
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # 设置列宽（每列约1.5英寸）
    for i in range(6):
        table.columns[i].width = Inches(1.5)
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '汉语翻译'
    header_cells[1].text = '英文单词'
    header_cells[2].text = '音标（答案）'
    header_cells[3].text = '汉语翻译'
    header_cells[4].text = '英文单词'
    header_cells[5].text = '音标（答案）'
    
    # 格式化表头
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.name = '宋体'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    
    # 填充数据（每行两个单词）
    for i in range(0, len(words_list), 2):
        row = table.add_row()
        
        # 第一个单词
        word1 = words_list[i]
        row.cells[0].text = word1['chinese']
        format_cell(row.cells[0], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        row.cells[1].text = '________________'
        format_cell(row.cells[1], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.underline = True
        
        row.cells[2].text = word1['phonetic']
        format_cell(row.cells[2], font_size=Pt(7), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for paragraph in row.cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(180, 180, 180)
        
        # 第二个单词（如果存在）
        if i + 1 < len(words_list):
            word2 = words_list[i + 1]
            row.cells[3].text = word2['chinese']
            format_cell(row.cells[3], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            row.cells[4].text = '________________'
            format_cell(row.cells[4], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for paragraph in row.cells[4].paragraphs:
                for run in paragraph.runs:
                    run.font.underline = True
            
            row.cells[5].text = word2['phonetic']
            format_cell(row.cells[5], font_size=Pt(7), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for paragraph in row.cells[5].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(180, 180, 180)
    
    doc.add_paragraph()  # 空行分隔

def generate_practice_english_to_chinese(doc, words_list, title="练习二：根据英文和音标填写汉语翻译"):
    """生成练习二：根据英文和音标填写汉语翻译"""
    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.size = Pt(14)
    
    # 添加说明
    instruction = doc.add_paragraph('说明：根据给出的英文单词和音标，在横线上填写对应的汉语翻译。')
    instruction.runs[0].font.size = Pt(10)
    instruction.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()  # 空行
    
    # 创建表格（6列：英文单词1、音标1、汉语翻译1、英文单词2、音标2、汉语翻译2）
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # 设置列宽（每列约1.5英寸）
    for i in range(6):
        table.columns[i].width = Inches(1.5)
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '英文单词'
    header_cells[1].text = '音标'
    header_cells[2].text = '汉语翻译'
    header_cells[3].text = '英文单词'
    header_cells[4].text = '音标'
    header_cells[5].text = '汉语翻译'
    
    # 格式化表头
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.name = '宋体'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    
    # 填充数据（每行两个单词）
    for i in range(0, len(words_list), 2):
        row = table.add_row()
        
        # 第一个单词
        word1 = words_list[i]
        row.cells[0].text = word1['word']
        format_cell(row.cells[0], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        row.cells[1].text = word1['phonetic']
        format_cell(row.cells[1], font_size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        row.cells[2].text = '________________'
        format_cell(row.cells[2], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for paragraph in row.cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.underline = True
        
        # 第二个单词（如果存在）
        if i + 1 < len(words_list):
            word2 = words_list[i + 1]
            row.cells[3].text = word2['word']
            format_cell(row.cells[3], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            row.cells[4].text = word2['phonetic']
            format_cell(row.cells[4], font_size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            row.cells[5].text = '________________'
            format_cell(row.cells[5], font_size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for paragraph in row.cells[5].paragraphs:
                for run in paragraph.runs:
                    run.font.underline = True

def shuffle_by_first_letter(words_list):
    """按首字母分组，组内打乱，然后按字母顺序合并"""
    # 按首字母分组
    groups = {}
    for word_data in words_list:
        # 获取首字母（大写），处理特殊情况如 "a/an"
        first_char = word_data['word'].strip()[0].upper()
        if first_char not in groups:
            groups[first_char] = []
        groups[first_char].append(word_data)
    
    # 在每个组内打乱
    for letter in groups:
        random.shuffle(groups[letter])
    
    # 按字母顺序合并所有组
    shuffled_words = []
    for letter in sorted(groups.keys()):
        shuffled_words.extend(groups[letter])
    
    return shuffled_words

def generate_practice_documents(json_file, output_dir='../out'):
    """生成两种默写练习文档"""
    # 加载单词
    words_list = load_words(json_file)
    
    # 按首字母分组并打乱
    shuffled_words = shuffle_by_first_letter(words_list)
    
    # 生成练习一：根据汉语翻译填写英文单词
    doc1 = Document()
    setup_document(doc1)
    
    # 添加总标题
    title1 = doc1.add_heading('KET 单词默写练习（一）', 0)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title1.runs[0].font.size = Pt(16)
    
    info1 = doc1.add_paragraph(f'共 {len(shuffled_words)} 个单词')
    info1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info1.runs[0].font.size = Pt(10)
    info1.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc1.add_paragraph()
    
    generate_practice_chinese_to_english(doc1, shuffled_words)
    
    output_file1 = f'{output_dir}/KET单词默写练习_中译英.docx'
    doc1.save(output_file1)
    print(f'成功生成练习一：{output_file1}')
    
    # 再次按首字母分组并打乱（生成不同的练习）
    shuffled_words = shuffle_by_first_letter(words_list)
    
    # 生成练习二：根据英文和音标填写汉语翻译
    doc2 = Document()
    setup_document(doc2)
    
    # 添加总标题
    title2 = doc2.add_heading('KET 单词默写练习（二）', 0)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.runs[0].font.size = Pt(16)
    
    info2 = doc2.add_paragraph(f'共 {len(shuffled_words)} 个单词')
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info2.runs[0].font.size = Pt(10)
    info2.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc2.add_paragraph()
    
    generate_practice_english_to_chinese(doc2, shuffled_words)
    
    output_file2 = f'{output_dir}/KET单词默写练习_英译中.docx'
    doc2.save(output_file2)
    print(f'成功生成练习二：{output_file2}')
    
    print(f'共包含 {len(words_list)} 个单词，已按首字母分组并在组内打乱')

if __name__ == '__main__':
    json_file = 'ket_A.json'
    generate_practice_documents(json_file)

