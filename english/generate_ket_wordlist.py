#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 KET 单词表 Word 文档
"""

import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def generate_ket_wordlist(json_file, output_file):
    """
    从 JSON 文件生成 KET 单词表 Word 文档
    
    Args:
        json_file: JSON 文件路径
        output_file: 输出的 Word 文档路径
    """
    # 读取 JSON 文件
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 提取所有单词
    words_list = []
    for letter, words in data.items():
        for word_data in words:
            words_list.append({
                'word': word_data.get('word', ''),
                'phonetic': word_data.get('phonetic', ''),
                'chinese': word_data.get('chinese', '')
            })
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置页面边距（更紧凑的布局）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # 添加标题
    title = doc.add_heading('KET 单词表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(14)  # 减小标题字体
    
    # 添加说明信息
    info_para = doc.add_paragraph(f'共 {len(words_list)} 个单词')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_format = info_para.runs[0].font
    info_format.size = Pt(9)
    info_format.color.rgb = RGBColor(128, 128, 128)
    info_para.paragraph_format.space_after = Pt(6)  # 减小间距
    
    # 创建表格（6列：单词1、音标1、翻译1、单词2、音标2、翻译2）
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # 设置表格列宽（每列约1.5英寸，总共约9英寸）
    for i in range(6):
        table.columns[i].width = Inches(1.5)
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '单词'
    header_cells[1].text = '音标'
    header_cells[2].text = '汉语翻译'
    header_cells[3].text = '单词'
    header_cells[4].text = '音标'
    header_cells[5].text = '汉语翻译'
    
    # 设置表头样式
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置中文字体
        cell.paragraphs[0].runs[0].font.name = '宋体'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置行高（更紧凑）
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    
    # 填充数据（每行两个单词）
    for i in range(0, len(words_list), 2):
        row = table.add_row()
        
        # 第一个单词
        word1 = words_list[i]
        row.cells[0].text = word1['word']
        row.cells[1].text = word1['phonetic']
        row.cells[2].text = word1['chinese']
        
        # 第二个单词（如果存在）
        if i + 1 < len(words_list):
            word2 = words_list[i + 1]
            row.cells[3].text = word2['word']
            row.cells[4].text = word2['phonetic']
            row.cells[5].text = word2['chinese']
        
        # 设置单元格样式
        for j, cell in enumerate(row.cells):
            # 设置字体大小（更小以容纳更多内容）
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    # 设置中文字体为宋体，英文字体为 Times New Roman
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                # 设置段落间距
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = Pt(12)  # 行距
            
            # 设置对齐方式
            col_index = j % 3
            if col_index == 0:  # 单词列：左对齐
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif col_index == 1:  # 音标列：居中
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:  # 翻译列：左对齐
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 保存文档
    doc.save(output_file)
    print(f'成功生成单词表：{output_file}')
    print(f'共包含 {len(words_list)} 个单词')

if __name__ == '__main__':
    json_file = 'ket_A.json'
    output_file = '../out/KET单词表.docx'
    generate_ket_wordlist(json_file, output_file)

