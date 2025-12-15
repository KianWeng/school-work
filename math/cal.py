import random
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def generate_addition_with_carry():
    """生成带进位的加法题目"""
    while True:
        a = random.randint(10, 99)
        b = random.randint(10, 99)
        # 确保有进位且结果不超过100
        if (a % 10) + (b % 10) >= 10 and a + b <= 100:
            return f"{a} + {b} ="

def generate_subtraction_with_borrow():
    """生成带退位的减法题目"""
    while True:
        a = random.randint(20, 99)
        b = random.randint(10, 99)
        # 确保有退位且结果为正数
        if a > b and (a % 10) < (b % 10):
            return f"{a} - {b} ="

def generate_multiplication():
    """生成9以内的乘法题目"""
    a = random.randint(1, 9)
    b = random.randint(1, 9)
    return f"{a} × {b} ="

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def create_single_page(doc, page_number):
    """在文档中添加一页内容"""
    # 添加标题
    title = doc.add_heading(f'口算练习题（第{page_number}页）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('姓名：___________    班级：___________    日期：___________')
    
    # 生成加减法题目（60道）
    doc.add_heading('一、加减法（60道）', level=1)
    
    # 生成30道加法
    addition_problems = []
    for _ in range(30):
        addition_problems.append(generate_addition_with_carry())
    
    # 生成30道减法
    subtraction_problems = []
    for _ in range(30):
        subtraction_problems.append(generate_subtraction_with_borrow())
    
    # 混合加减法题目
    all_add_sub = addition_problems + subtraction_problems
    random.shuffle(all_add_sub)
    
    # 将题目分成5列显示
    for i in range(0, len(all_add_sub), 5):
        row_problems = all_add_sub[i:i+5]
        # 创建表格行
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        for j, problem in enumerate(row_problems):
            cell = table.cell(0, j)
            cell.text = problem
            # 设置字体大小
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        # 如果这一行不足5个题目，填充空单元格
        for j in range(len(row_problems), 5):
            cell = table.cell(0, j)
            cell.text = ""
    
    # 添加乘法题目（40道）
    doc.add_heading('二、乘法（40道）', level=1)
    
    multiplication_problems = []
    for _ in range(40):
        multiplication_problems.append(generate_multiplication())
    
    # 将乘法题目分成5列显示
    for i in range(0, len(multiplication_problems), 5):
        row_problems = multiplication_problems[i:i+5]
        # 创建表格行
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        for j, problem in enumerate(row_problems):
            cell = table.cell(0, j)
            cell.text = problem
            # 设置字体大小
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        # 如果这一行不足5个题目，填充空单元格
        for j in range(len(row_problems), 5):
            cell = table.cell(0, j)
            cell.text = ""

def create_word_document_with_pages(num_pages=10):
    """创建包含多页的Word文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 添加总标题
    main_title = doc.add_heading('口算练习题集', 0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分页符
    add_page_break(doc)
    
    # 生成每一页
    for page_num in range(1, num_pages + 1):
        print(f"正在生成第{page_num}页...")
        create_single_page(doc, page_num)
        
        # 如果不是最后一页，添加分页符
        if page_num < num_pages:
            add_page_break(doc)
    
    # 保存文档
    filename = f'口算练习题_{num_pages}页.docx'
    doc.save(filename)
    print(f"\n✅ 口算练习题已生成并保存为 '{filename}'")
    print(f"文档包含{num_pages}页，每页100道题目")
    return filename

if __name__ == "__main__":
    create_word_document_with_pages(10)
