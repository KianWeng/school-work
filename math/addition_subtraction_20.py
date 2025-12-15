import random
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_addition():
    """生成20以内的加法题目"""
    while True:
        a = random.randint(1, 19)
        b = random.randint(1, 19)
        # 结果不超过20
        if a + b <= 20:
            return f"{a} + {b} ="

def generate_subtraction():
    """生成20以内的减法题目"""
    while True:
        a = random.randint(1, 20)
        b = random.randint(1, 20)
        # 被减数大于等于减数，结果非负
        if a >= b:
            return f"{a} - {b} ="

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def add_problems_to_document(doc, problems, title, heading_level=1):
    """将题目添加到文档中"""
    doc.add_heading(title, level=heading_level)
    
    # 将题目分成5列显示
    for i in range(0, len(problems), 5):
        row_problems = problems[i:i+5]
        # 创建表格行
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        for j, problem in enumerate(row_problems):
            cell = table.cell(0, j)
            cell.text = problem
            # 设置字体大小
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        # 如果这一行不足5个题目，填充空单元格
        for j in range(len(row_problems), 5):
            cell = table.cell(0, j)
            cell.text = ""

def create_single_page(doc, page_number):
    """在文档中添加一页内容"""
    # 添加标题
    title = doc.add_heading(f'20以内加减法练习题（第{page_number}页）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('日期：___________    时间：___________    得分：___________')
    
    # 生成加法题目（50道）
    addition_problems = []
    for _ in range(50):
        addition_problems.append(generate_addition())
    
    # 生成减法题目（50道）
    subtraction_problems = []
    for _ in range(50):
        subtraction_problems.append(generate_subtraction())
    
    # 合并所有题目并打乱顺序
    all_problems = addition_problems + subtraction_problems
    random.shuffle(all_problems)
    
    # 添加总标题并显示所有题目
    add_problems_to_document(doc, all_problems, '口算练习题（100道）')

def create_word_document_with_pages(num_pages=10):
    """创建包含多页的Word文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 添加总标题
    main_title = doc.add_heading('20以内加减法练习题集', 0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分页符
    add_page_break(doc)
    
    # 生成每一页
    for page_num in range(1, num_pages + 1):
        print(f"正在生成第{page_num}页...")
        create_single_page(doc, page_num)
        
        # 如果不是最后一页，添加分页符
        if page_num < num_pages:
            add_page_break(doc)
    
    # 保存文档
    filename = f'20以内加减法练习题_{num_pages}页.docx'
    doc.save(filename)
    print(f"\n✅ 20以内加减法练习题已生成并保存为 '{filename}'")
    print(f"文档包含{num_pages}页，每页100道题目（50道加法 + 50道减法，顺序随机）")
    return filename

if __name__ == "__main__":
    create_word_document_with_pages(50)

